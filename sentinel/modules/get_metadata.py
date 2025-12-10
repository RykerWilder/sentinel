import os
import sys
from colorama import Style, Fore
from pathlib import Path
from datetime import datetime, timedelta
from sentinel import write_to_result_file
import mimetypes

class ExtractMetaData:
    def get_basic_metadata(self, filepath):
        stats = os.stat(filepath)
        return {
            'file_name': os.path.basename(filepath),
            'size': self.format_size(stats.st_size),
            'creation_date': datetime.fromtimestamp(stats.st_ctime).isoformat(),
            'edit_date': datetime.fromtimestamp(stats.st_mtime).isoformat(),
            'access_date': datetime.fromtimestamp(stats.st_atime).isoformat(),
            'mime_type': mimetypes.guess_type(filepath)[0] or 'unknown'
        }

    def format_size(self, bytes_size):
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if bytes_size < 1024.0:
                return f"{bytes_size:.2f} {unit}"
            bytes_size /= 1024.0
        return f"{bytes_size:.2f} PB"

    def extract_image_metadata(self, filepath):
        try:
            from PIL import Image
            from PIL.ExifTags import TAGS
            
            img = Image.open(filepath)
            metadata = {
                'format': img.format,
                'mode': img.mode,
                'size': f"{img.width}x{img.height}",
                'width': img.width,
                'height': img.height
            }
            
            #EXTRACT EXIF DATA IF EXISTS
            exif_data = img._getexif()
            if exif_data:
                metadata['exif'] = {}
                for tag_id, value in exif_data.items():
                    tag = TAGS.get(tag_id, tag_id)
                    metadata['exif'][tag] = str(value)
            
            return metadata
        except Exception as e:
            return {'error': str(e)}

    def extract_video_metadata(self, filepath):
        try:
            from pymediainfo import MediaInfo
            
            media_info = MediaInfo.parse(filepath)
            metadata = {}
            
            for track in media_info.tracks:
                if track.track_type == 'General':
                    metadata['duration'] = str(timedelta(milliseconds=track.duration)) if track.duration else None
                    metadata['format'] = track.format
                    metadata['total_bitrate'] = track.overall_bit_rate
                elif track.track_type == 'Video':
                    metadata['video_codec'] = track.codec
                    metadata['resolution'] = f"{track.width}x{track.height}" if track.width else None
                    metadata['fps'] = track.frame_rate
                    metadata['bitrate'] = track.bit_rate
                elif track.track_type == 'Audio':
                    metadata['codec_audio'] = track.codec
                    metadata['audio_channels'] = track.channel_s
                    metadata['sample_rate'] = track.sampling_rate
                    metadata['bitrate_audio'] = track.bit_rate
            
            return metadata
        except Exception as e:
            return {'error': str(e)}

    def extract_audio_metadata(self, filepath):
        try:
            from mutagen import File
            
            audio = File(filepath, easy=True)
            if audio is None:
                return {'Error': 'Audio file not recognized'}
            
            metadata = {
                'duration': f"{audio.info.length:.2f} seconds" if hasattr(audio, 'info') else None,
                'bitrate': f"{audio.info.bitrate} bps" if hasattr(audio, 'info') and hasattr(audio.info, 'bitrate') else None,
                'sample_rate': f"{audio.info.sample_rate} Hz" if hasattr(audio, 'info') and hasattr(audio.info, 'sample_rate') else None
            }
            
            #COMMON TAGS
            if audio.tags:
                metadata['tags'] = dict(audio.tags)
            
            return metadata
        except Exception as e:
            return {'error': str(e)}

    def extract_pdf_metadata(self, filepath):
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(filepath)
            metadata = {
                'pages_number': len(reader.pages),
                'document_info': {}
            }
            
            if reader.metadata:
                for key, value in reader.metadata.items():
                    metadata['document_info'][key] = value
            
            return metadata
        except Exception as e:
            return {'error': str(e)}

    def extract_docx_metadata(self, filepath):
        try:
            from docx import Document
            
            doc = Document(filepath)
            core_props = doc.core_properties
            
            metadata = {
                'author': core_props.author,
                'title': core_props.title,
                'subject': core_props.subject,
                'creation_date': core_props.created.isoformat() if core_props.created else None,
                'modification_date': core_props.modified.isoformat() if core_props.modified else None,
                'last_edit': core_props.last_modified_by,
                'paragraphs_number': len(doc.paragraphs),
                'tables_number': len(doc.tables)
            }
            
            return metadata
        except Exception as e:
            return {'error': str(e)}

    def extract_metadata(self, filepath):
        if not os.path.exists(filepath):
            return {'error': 'File not found'}
        
        if not os.path.isfile(filepath):
            return {'error': 'The path does not point to a file'}
        
        metadata = {'basic_metadata': self.get_basic_metadata(filepath)}
        
        mime_type = metadata['basic_metadata']['mime_type']
        ext = os.path.splitext(filepath)[1].lower()
        
        if mime_type and mime_type.startswith('image'):
            metadata['image_metadata'] = self.extract_image_metadata(filepath)
        elif mime_type and mime_type.startswith('video'):
            metadata['video_metadata'] = self.extract_video_metadata(filepath)
        elif mime_type and mime_type.startswith('audio'):
            metadata['audio_metadata'] = self.extract_audio_metadata(filepath)
        elif ext == '.pdf':
            metadata['pdf_metadata'] = self.extract_pdf_metadata(filepath)
        elif ext in ['.docx', '.doc']:
            metadata['document_metadata'] = self.extract_docx_metadata(filepath)
        else:
            metadata['note'] = 'File type not supported for specific metadata'
        
        return metadata

    def format_metadata_to_string(self, metadata, indent=0):
        """Converte i metadata in una stringa formattata"""
        result = []
        spacing = "  " * indent
        for key, value in metadata.items():
            if isinstance(value, dict):
                result.append(f"{spacing}{key}:")
                result.append(self.format_metadata_to_string(value, indent + 1))
            else:
                result.append(f"{spacing}{key}: {value}")
        return "\n".join(result)

    def get_metadata_manager(self):
        print(f"\n{'='*40}{Fore.MAGENTA} METADATA EXTRACTOR {Style.RESET_ALL}{'='*40}")
        
        filepath = input(f"\n{Fore.CYAN}┌─[Enter the absolute path to the file] \n└──> {Style.RESET_ALL}").strip()
        
        filepath = filepath.strip('"').strip("'")
        
        metadata = self.extract_metadata(filepath)
        
        #CONVERT METADAT IN STRING
        metadata_string = self.format_metadata_to_string(metadata)

        content = f"\nMETADATA EXTRACTION RESULTS"
        content += f"File path: {filepath}\n"
        content += metadata_string
        
        #SAVE ON FILE
        result_file = write_to_result_file(content, "MetadataExtractor")
        
        if result_file:
            print(f"\n{Fore.MAGENTA}[INFO]{Style.RESET_ALL} Metadata extracted successfully and saved to: {result_file}")
        else:
            print(f"\n{Fore.RED}[X] Error saving metadata to file{Style.RESET_ALL}")